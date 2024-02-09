from flask import Flask, request, render_template
from docx import Document
from num2words import num2words
from datetime import datetime, timedelta

app = Flask(__name__)

def valor_por_extenso(valor):
    valor_extenso = num2words(valor, lang='pt_BR', to='currency')
    return valor_extenso.capitalize()

@app.route('/')
def formulario_contrato():
    return render_template('formulario_contrato.html')

@app.route('/gerar_contrato', methods=['POST'])
def gerar_contrato():
    nome_locatario = request.form['nome_locatario']
    cpf_locatario = request.form['cpf_locatario']
    telefone_locatario = request.form['telefone_locatario']
    valor_aluguel = float(request.form['valor_aluguel'])
    valor_calcao = valor_aluguel * 2

    doc = Document()
    
    # Lista de meses em português
    meses = ['janeiro', 'fevereiro', 'março', 'abril', 'maio', 'junho', 'julho', 'agosto', 'setembro', 'outubro', 'novembro', 'dezembro']

    # Data atual
    data_atual = datetime.now()
    data_termino = data_atual + timedelta(days=365)  # Calcula a data de término como 1 ano após a data atual

    # Adicionando o texto do contrato centralizado
    p_contrato = doc.add_paragraph("CONTRATO DE LOCAÇÃO DE IMÓVEL\n\n")
    p_contrato.alignment = 1  # Define o alinhamento do parágrafo como centralizado
    
    p_locador = doc.add_paragraph()
    p_locador.add_run("LOCADOR:").bold = True
    p_locador.add_run(" PP. Tereza Maria Cardoso da Silva, brasileira, inscrita no CPF 31589413091 e RG 1008298679 residente e domiciliado nesta capital, AV. Bernardi n° 437 apto 202, Bairro Cristo Redentor, CEP 91040-030, FONE 999414337\n")
    
    p_locatario = doc.add_paragraph()
    p_locatario.add_run("LOCATÁRIO: ").bold = True
    p_locatario.add_run(f"{nome_locatario}, brasileiro(a), portador(a) do CPF {cpf_locatario}, Fone {telefone_locatario}")

    p_imovel = doc.add_paragraph()
    p_imovel.add_run("IMÓVEL: ").bold = True
    p_imovel.add_run("AV. Brino 371, apto 102, Santa Maria Goretti, Porto Alegre, RS.\n\n")

    doc.add_paragraph("As partes acima qualificadas, pelo presente instrumento, contratam a locação do imóvel supracitado, mediante as clausulas e condições seguintes:\n\n")

    p_primeira = doc.add_paragraph()
    p_primeira.add_run("PRIMEIRA: ").bold = True
    p_primeira.add_run("O prazo de locação é de 1(um) ano, iniciando em ")
    p_primeira.add_run(f"{data_atual.day} de {meses[data_atual.month - 1]} de {data_atual.year}")  # Formata a data atual no formato desejado
    p_primeira.add_run(f", terminado de pleno direito no dia {data_termino.day} de {meses[data_termino.month - 1]} de {data_termino.year}, independente de notificação ou aviso judicial ou extrajudicial.\n\n")
    doc.add_paragraph("Parágrafo Único: Estando a locação por prazo indeterminado, o contrato somente poderá ser resilido mediante prévio aviso de e 30 dias, independente de notificação ou aviso judicial ou extrajudicial 30 (trinta) dias, por escrito, sob pena de pagamento de 01 (um) mês de aluguel e encargos vigentes a época de resilição.\n\n")

    p_segunda = doc.add_paragraph()
    p_segunda.add_run("SEGUNDA: ").bold = True
    p_segunda.add_run(f"O aluguel inicial é de R${valor_aluguel:.2f} ({valor_por_extenso(valor_aluguel)}) mais taxas, a ser efetuado diretamente ao locador. Devendo fazê-lo até o quinto dia útil de cada mês, subseqüente ao vencido, sob pena de multa e correções, acrescidos de multa de 10% (dez por cento) ao mês, juros moratórios de 1% (um) ao mês e correção monetária.\n\n")

    p_terceira = doc.add_paragraph()
    p_terceira.add_run("TERCEIRA: ").bold = True
    p_terceira.add_run(f"O locatário concorda desde já em depositar a titulo de fiança a caução no valor de R${valor_calcao:.2f} ({valor_por_extenso(valor_calcao)}) equivalentes a 02 (dois) meses de aluguel.\n\n")

    p_quarta = doc.add_paragraph()
    p_quarta.add_run("QUARTA: ").bold = True
    p_quarta.add_run("O locatário obriga-se a zelar pela conservação do imóvel e a fazer de imediato e por sua conta todas as reparações dos estragos a que der causa no curso da locação, de modo especial as referentes a vazamentos e obstruções que venham a surgir no sistema de água e esgoto, devendo restituir o imóvel, no fim da locação, no mesmo estado em que recebeu, salvo as deteriorações decorrentes do uso normal e ainda com pintura nova. No final da locação o Locatário obriga-se a refazer a pintura, de acordo com a vistoria, sendo que o material a ser empregado deverá ser idêntico ao material do início da locação. Ocorrendo infiltração de água de/ou para imóvel vizinho, o fato deverá ser comunicado por escrito, de imediato ao locador para as providências cabíveis.\n\n")

    p_quinta = doc.add_paragraph()
    p_quinta.add_run("QUINTA: ").bold = True
    p_quinta.add_run("Todas as despesas diretamente ligadas a conservação do imóvel, tais como água, luz, gás, telefone ficarão sob responsabilidade do locatário.\n\n")

    p_recebi = doc.add_paragraph()
    p_recebi.add_run(f"Recebi a quantia de R${valor_calcao:.2f} ({valor_por_extenso(valor_calcao)}) referente a dois meses de caução deste contrato.\n\n").bold = True

    doc.add_paragraph("E por estarem justos e contratados lavraram instrumento em 02 (duas) vias de igual teor e forma para as finalidades de direito.\n\n")

    doc.add_paragraph("Porto Alegre, ")
    p_data = doc.add_paragraph()
    p_data.add_run(f"{data_atual.day} de {meses[data_atual.month - 1]} de {data_atual.year}").bold = True
    p_data.add_run("\n\n\n\n\n\n")

    p_assinatura = doc.add_paragraph()
    p_assinatura.add_run("_________________________________________" + " " * 30 + "________________________________________")
    doc.add_paragraph("PP: Tereza Maria Cardoso da Silva" + " " * 30 + f"Locatário: {nome_locatario}\n\n")

    # Salvar o documento com o nome do locatário (sem espaços)
    nome_arquivo = nome_locatario.replace(" ", "_")
    doc.save(f"contrato_de_locacao_{nome_arquivo}.docx")

    return f"Contrato gerado com sucesso para {nome_locatario}. Verifique o arquivo 'contrato_de_locacao_{nome_arquivo}.docx'."

if __name__ == '__main__':
    app.run(debug=True, host='0.0.0.0', port=5000)

